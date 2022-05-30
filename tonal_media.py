import asyncio
import re
from datetime import datetime
from datetime import timedelta
import logging.config
import dateutil
import docx
import httpx
from docx.oxml.shared import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from pptx.dml.color import RGBColor

from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from pptx.chart.data import CategoryChartData
from pptx.enum.chart import XL_CHART_TYPE

from resp import post
from word_media import add_title_data, login
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger('foo-logger')

CLEANR = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')

TIMEOUT = 15 * 60
WEEK_TRUST = "https://api.glassen-it.com/component/socparser/content/getweektrust"
THREAD_STATS = "https://api.glassen-it.com/component/socparser/stats/getThreadStats"

THREAD_DATA = "https://api.glassen-it.com/component/socparser/thread/additional_info"
OWNERS_TOP = "https://api.glassen-it.com/component/socparser/stats/getOwnersTopByPostCount"

SUBECT_URL = "https://api.glassen-it.com/component/socparser/users/getreferences"
STATISTIC_POST_URL = "https://api.glassen-it.com/component/socparser/content/posts"
STYLE = "Times New Roman"
PT = Pt(10.5)


async def docx_tonal(thread_ids, _from, _to, iogv_name, types, smi_type):
    from app import add_title_text
    response_data = [await get_session_tonal_result(thread_id, _from, _to, types, smi_type) for thread_id in thread_ids]
    name = []
    for thread_, _, _, _ in response_data:
        name.append(thread_['name'])

    document = Document()
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style(STYLE, WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.name = STYLE
    obj_font.size = Pt(16)

    _from_str = dateutil.parser.parse(_from).strftime('%Y-%m-%d')
    _to_str = dateutil.parser.parse(_to).strftime('%Y-%m-%d')

    named = ", ".join(name)
    add_title_text(document,
                   f'ИОГВ: {iogv_name}, отчет по лентам {named} за период с {_from_str} по {_to_str}',
                   False
                   )
    table_number = 1
    flot_number = 1
    i = 0
    for thread_name, week_trust, owners_top_smi, owners_top_social in response_data:
        i +=1
        if i > 1:
            document.add_page_break()

        parag_title = document.add_paragraph()

        title_run = parag_title.add_run(name, style=STYLE)
        title_run.bold = True
        title_run.font.size = Pt(12)

        add_table_title = True
        add_table_general(document, table_number, "Показатели", week_trust, add_table_title)
        if "smi" in types:
            table_number += 1
            add_table_owners(document, table_number, "Топ Источников СМИ ", owners_top_smi, add_table_title)
        if "social" in types:
            table_number += 1
            add_table_owners(document, table_number, "Топ Источников Соц. Сети ", owners_top_social, add_table_title)
        if ("smi" in types and len(owners_top_smi['items']) > 0) or ("social" in types and len(owners_top_social['items']) > 0):
            document.add_page_break()
        if "smi" in types:
            flot_number = add_table_tonal(document, "СМИ", flot_number, owners_top_smi)
        if "social" in types:
            flot_number = add_table_tonal(document, "Соц. Сети", flot_number, owners_top_social)
    return document


async def get_session_tonal_result(thread_id, _from, _to, types, smi_type):
    async with httpx.AsyncClient() as session:
        session = await login(session)
        try:
            await get_thread_name(session, thread_id)
        except Exception as e:
            print(e)

        thread_name, week_trust_smi, week_trust_social, owners_top_smi, owners_top_social = await asyncio.gather(
            get_thread_name(session, thread_id),
            get_stats_trust(session, thread_id, _from, _to, types, smi_type),
            get_stats_trust(session, thread_id, _from, _to, types, smi_type, "social"),
            get_owners_top(session, thread_id, _from, _to, types, smi_type),
            get_owners_top(session, thread_id, _from, _to, types, smi_type, "social"),

        )
        week_trust = {}
        try:
            publication_positive_social = week_trust_social['stats']['positive']
            publication_netural_social = week_trust_social['stats']['netural']
            publication_negative_social = week_trust_social['stats']['negative']
        except Exception as e:
            publication_positive_social = {'posts_count': 0, 'comments_count': 0, 'reposts_count': 0, 'viewed_count': 0,
                                           'likes_count': 0}

            publication_netural_social = {'posts_count': 0, 'comments_count': 0, 'reposts_count': 0, 'viewed_count': 0,
                                          'likes_count': 0}

            publication_negative_social = {'posts_count': 0, 'comments_count': 0, 'reposts_count': 0, 'viewed_count': 0,
                                           'likes_count': 0}

        try:
            publication_positive_smi = week_trust_smi['stats']['positive']
            publication_netural_smi = week_trust_smi['stats']['netural']
            publication_negative_smi = week_trust_smi['stats']['negative']
        except Exception as e:
            publication_positive_smi = {'posts_count': 0, 'comments_count': 0, 'reposts_count': 0, 'viewed_count': 0,
                                        'likes_count': 0}
            publication_netural_smi = {'posts_count': 0, 'comments_count': 0, 'reposts_count': 0, 'viewed_count': 0,
                                       'likes_count': 0}
            publication_negative_smi = {'posts_count': 0, 'comments_count': 0, 'reposts_count': 0, 'viewed_count': 0,
                                        'likes_count': 0}

        week_trust["Количество публикаций"] = [
            publication_positive_social['posts_count'] + publication_positive_smi['posts_count']
            + publication_netural_social['posts_count'] + publication_netural_smi['posts_count'] +
            publication_negative_social['posts_count'] + publication_negative_smi['posts_count'],
            publication_positive_social['posts_count'] + publication_positive_smi['posts_count'],
            publication_netural_social['posts_count'] + publication_netural_smi['posts_count'],
            publication_negative_social['posts_count'] + publication_negative_smi['posts_count'],

        ]

        week_trust["Количество комментариев"] = [
            publication_positive_social['comments_count'] + publication_positive_smi['comments_count']
            + publication_netural_social['comments_count'] + publication_netural_smi['comments_count'] +
            publication_negative_social['comments_count'] + publication_negative_smi['comments_count'],
            publication_positive_social['comments_count'] + publication_positive_smi['comments_count'],
            publication_netural_social['comments_count'] + publication_netural_smi['comments_count'],
            publication_negative_social['comments_count'] + publication_negative_smi['comments_count'],
        ]
        week_trust["Количество репостов"] = [
            publication_positive_social['reposts_count'] + publication_positive_smi['reposts_count']
            + publication_netural_social['reposts_count'] + publication_netural_smi['reposts_count'] +
            publication_negative_social['reposts_count'] + publication_negative_smi['reposts_count'],
            publication_positive_social['reposts_count'] + publication_positive_smi['reposts_count'],
            publication_netural_social['reposts_count'] + publication_netural_smi['reposts_count'],
            publication_negative_social['reposts_count'] + publication_negative_smi['reposts_count'],
        ]
        week_trust["Количество лайков"] = [
            publication_positive_social['likes_count'] + publication_positive_smi['likes_count']
            + publication_netural_social['likes_count'] + publication_netural_smi['likes_count'] +
            publication_negative_social['likes_count'] + publication_negative_smi['likes_count'],
            publication_positive_social['likes_count'] + publication_positive_smi['likes_count'],
            publication_netural_social['likes_count'] + publication_netural_smi['likes_count'],
            publication_negative_social['likes_count'] + publication_negative_smi['likes_count'],
        ]
        week_trust["Количество просмотров"] = [
            publication_positive_social['viewed_count'] + publication_positive_smi['viewed_count']
            + publication_netural_social['viewed_count'] + publication_netural_smi['viewed_count'] +
            publication_negative_social['viewed_count'] + publication_negative_smi['viewed_count'],
            publication_positive_social['viewed_count'] + publication_positive_smi['viewed_count'],
            publication_netural_social['viewed_count'] + publication_netural_smi['viewed_count'],
            publication_negative_social['viewed_count'] + publication_negative_smi['viewed_count'],
        ]
        return thread_name, week_trust, owners_top_smi, owners_top_social


async def get_thread_name(session, thread_id):
    response = await post(session, THREAD_DATA, {
        "thread_id": thread_id,
    })
    response_json = response.json()
    return response_json


async def get_owners_top(session, thread_id, _from, _to, types, smi_type, type="smi"):
    if type not in types:
        return []
    response = await post(session, OWNERS_TOP, {
        "thread_id": thread_id,
        "from": _from,
        "to": _to,
        "type": type,
        "smi_type": smi_type

    })
    response_json = response.json()
    return response_json


async def get_stats_trust(session, thread_id, _from, _to, types, smi_type, type="smi"):
    if type not in types:
        return []
    payload = {"thread_id": thread_id,
               "from": _from,
               "to": _to,
               "type": type,
                "smi_type": smi_type
               }
    response = await post(session, THREAD_STATS, payload)
    response_json = response.json()
    return response_json


async def get_week_trust(session, thread_id, _from):
    from app import UTC
    days = (datetime.today() + timedelta(hours=UTC) - dateutil.parser.parse(_from)).days

    payload = {
        "thread_id": thread_id,
        "days": days
    }
    response = await post(session, WEEK_TRUST, payload)
    response_json = response.json()
    result = {}
    result["Количество публикаций"] = [
        response_json['trust'][0]['positive'] + response_json['trust'][0]['netural'] + response_json['trust'][0][
            'negative'],
        response_json['trust'][0]['positive'],
        response_json['trust'][0]['netural'],
        response_json['trust'][0]['negative']
    ]
    positive = response_json['trust_extended']['1']['positive']
    netural = response_json['trust_extended']['1']['netural']
    negative = response_json['trust_extended']['1']['negative']

    result["Количество комментариев"] = [
        positive['comments'] + netural['comments'] + negative['comments'],
        positive['comments'],
        netural['comments'],
        negative['comments']
    ]
    result["Количество репостов"] = [
        positive['reposts'] + netural['reposts'] + negative['reposts'],
        positive['reposts'],
        netural['reposts'],
        negative['reposts']
    ]
    result["Количество лайков"] = [
        positive['likes'] + netural['likes'] + negative['likes'],
        positive['likes'],
        netural['likes'],
        negative['likes']
    ]
    result["Количество просмотров"] = [
        positive['viewed'] + netural['viewed'] + negative['viewed'],
        positive['viewed'],
        netural['viewed'],
        negative['viewed']
    ]
    return result


def add_table_general(document, table_number, header, records, add_table_title):
    from app import add_name
    from app import set_center, set_right
    from app import set_cell_vertical_alignment

    # if add_table_title:
    #     parag_table = document.add_paragraph()
    #     parag_table.add_run(
    #         "",
    #         style=STYLE
    #     )
    #     parag_table.style.font.size = Pt(1)

    parag_table_1 = document.add_paragraph()

    text = f' Таблица {table_number} - {header} '

    parag_table_1.add_run(
        text,
        style=STYLE
    )

    parag_table_1.paragraph_format.space_after = Inches(0)
    parag_table_1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    parag_table_1.runs[-1].font.size = Pt(15)

    table = document.add_table(rows=0, cols=5)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(2.9)
    table.columns[1].width = Inches(0.9)
    table.columns[2].width = Inches(0.9)
    table.columns[3].width = Inches(0.9)
    table.columns[4].width = Inches(0.9)

    table.style = 'TableGrid'

    hdr_cells = table.add_row().cells

    hdr_cells[0].text = "Название"
    hdr_cells[1].text = "Всего"
    hdr_cells[2].text = "Поз."
    hdr_cells[3].text = "Нейт."
    hdr_cells[4].text = "Негат."

    set_cell_vertical_alignment(hdr_cells[1])

    for k, v in records.items():
        row_cells = table.add_row().cells
        row_cells[4].text = str(v[3])
        row_cells[3].text = str(v[2])
        row_cells[2].text = str(v[1])

        row_cells[1].text = str(v[0])
        row_cells[0].text = str(k)
        row_cells[0].alignment = WD_TABLE_ALIGNMENT.LEFT
        row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_right(row_cells[1])
        set_right(row_cells[2])
        set_right(row_cells[3])
        set_right(row_cells[4])

        # i += 1


def add_table_owners(document, table_number, header, records, add_table_title, type="СМИ"):
    from app import add_name
    from app import set_center, set_right
    from app import set_cell_vertical_alignment
    from app import add_hyperlink

    if add_table_title:
        parag_table = document.add_paragraph()
        parag_table.add_run(
            "",
            style=STYLE
        )
        parag_table.style.font.size = Pt(1)

    parag_table_1 = document.add_paragraph()

    text = f' Таблица {table_number} - {header} '

    parag_table_1.add_run(
        text,
        style=STYLE
    )
    parag_table_1.runs[-1].font.size = Pt(15)
    parag_table_1.paragraph_format.space_after = Inches(0)
    parag_table_1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    table = document.add_table(rows=0, cols=6)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(2.1)
    table.columns[2].width = Inches(0.6)
    table.columns[3].width = Inches(0.6)
    table.columns[4].width = Inches(0.6)
    table.columns[5].width = Inches(0.6)

    table.style = 'TableGrid'

    hdr_cells = table.add_row().cells

    hdr_cells[0].text = "Название"
    hdr_cells[1].text = "URL"
    hdr_cells[2].text = "Всего"
    hdr_cells[3].text = "Поз."
    hdr_cells[4].text = "Нейт."
    hdr_cells[5].text = "Негат."

    set_cell_vertical_alignment(hdr_cells[1])

    for d in records.get('items', []):
        row_cells = table.add_row().cells
        row_cells[5].text = str(d.get('negative'))
        row_cells[4].text = str(d.get('netural'))
        row_cells[3].text = str(d.get('positive'))
        row_cells[2].text = str(d.get('post_count'))
        row_cells[1].text = ""

        row_cells[0].text = str(d.get('title'))
        row_cells[0].alignment = WD_TABLE_ALIGNMENT.LEFT
        row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_center(row_cells[1])
        set_right(row_cells[2])
        set_right(row_cells[3])
        set_right(row_cells[4])
        set_right(row_cells[5])

        row_cells[1].paragraphs[0].runs[-1].italic = True
        row_cells[1].paragraphs[0].runs[-1].bold = True
        row_cells[1].paragraphs[0].runs[-1].font.size = Pt(8)
        # row_cells[1].paragraphs[0].add_run(
        #     "",
        #     style=STYLE
        # )
        add_hyperlink(row_cells[1].paragraphs[0], str(d.get('url')), str(d.get('url')), None, True)


def add_table_tonal(document, type, chart_number, data_soc):
    categories_data = {}
    categories_data_comments = {}
    data = []
    if len(data_soc['items']) <= 0:
        return chart_number
    for items in data_soc['items']:
        data.extend(items['graph'])
    if len(data) > 0:
        sorted(data, key=lambda x: dateutil.parser.parse(x['item_date']))
    else:
        return chart_number

    if len(data[0]['item_date']) > 10:
        first_date = dateutil.parser.parse(data[0]['item_date']).hour
        last_date = dateutil.parser.parse(data[-1]['item_date']).hour
        if first_date < last_date:
            categories = list(range(first_date, last_date + 1))
        else:
            categories = list(range(first_date, 24)) + list(range(0, last_date + 1))
        for c in categories:
            categories_data[c] = 0
            categories_data_comments[c] = 0
        for d in data:
            hour = dateutil.parser.parse(d['item_date']).hour
            categories_data[hour] += int(d["post_count"])
            categories_data_comments[hour] += int(d["comments_count"])

        categories_list = [f"{x}:00" for x in categories]

    else:
        first_date = dateutil.parser.parse(data[0]['item_date'])
        last_date = dateutil.parser.parse(data[-1]['item_date'])
        categories = [first_date]
        while_first_date = first_date
        while True:
            while_first_date += timedelta(days=1)
            if while_first_date <= last_date:
                categories.append(while_first_date)
            else:
                break
        for c in categories:
            categories_data[c] = 0
            categories_data_comments[c] = 0
        for d in data:
            hour = dateutil.parser.parse(d['item_date'])
            categories_data[hour] += int(d["post_count"])
            categories_data_comments[hour] += int(d["comments_count"])
        categories_list = [f"{x.day}.{x.month}" for x in categories]

    from app import change_color, update_chart_style
    parag_table = document.add_paragraph()
    parag_table.add_run(
        f' График {chart_number} - Динамика распространения публикаций по топ {type}',
        style=STYLE
    )

    parag_table.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    chart_data = CategoryChartData()
    chart_data.categories = categories_list



    chart_data.add_series('Публикации', list(categories_data.values()))
    chart_data.add_series('Комментарии', list(categories_data_comments.values()))

    x, y, cx, cy = Inches(-3.5), Inches(0), Inches(6.15), Inches(3.3)

    chart = document.add_chart(XL_CHART_TYPE.COLUMN_CLUSTERED, x, y, cx, cy, chart_data)
    change_color(chart.plots[0].series[0], RGBColor(114, 159, 207))
    change_color(chart.plots[0].series[1], RGBColor(87, 57, 132))



    chart_number += 1
    update_chart_style(chart)
    return chart_number
