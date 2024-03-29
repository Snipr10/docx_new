import asyncio
import re
from datetime import datetime
from datetime import timedelta
import logging.config
import dateutil
import docx
import httpx
from docx.oxml.shared import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor

from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn



from resp import post
from settings import LOGIN_URL, SUBECT_URL, STATISTIC_POST_URL, login_l, password_p, NETWORK_IDS


CLEANR = re.compile('<.*?>|&([a-z0-9]+|#[0-9]{1,6}|#x[0-9a-f]{1,6});')

TIMEOUT = 15 * 60

STYLE = "Times New Roman"
PT = Pt(10.5)
DATE_FORMAT = "%d-%m-%Y %H:%M:%S"


async def login(session, login=login_l, password=password_p):
    try:
        from app import logger
    except Exception:
        logger = logging.getLogger('foo-logger')
    uid = None
    try:
        from app import COOKIES
        # if len(COOKIES) == 0 or COOKIES[0].get("date") < datetime.today() - timedelta(minutes=10):
        #     payload = {
        #         "login": login,
        #         "password": password
        #     }
        #     response = await post(session, LOGIN_URL, payload)
        #
        #     if response.status_code != 200:
        #         logger.error(f"login {response}")
        #
        #         raise Exception("can not login")
        #     COOKIES.clear()
        #     COOKIES.append({"date": datetime.now(), "cookies": session.cookies})
        # else:
        #     session.cookies = COOKIES[0].get("cookies")
        payload = {
            "login": login,
            "password": password
        }
        response = await post(session, LOGIN_URL, payload)
        logger.error(f"login {response}")
        logger.error(f"login {response.status_code}")
        logger.error(f"login {response.text}")

        if response.status_code != 200:
            logger.error(f"login {response}")
            raise Exception("can not login")

    except Exception as e:
        logger.error(f"login 1 {e}")

        payload = {
            "login": login,
            "password": password
        }
        response = await post(session, LOGIN_URL, payload)

        if response.status_code != 200:
            logger.error(f"login {response}")
            raise Exception("can not login")
        logger.error(f"login success {login}")
        uid = response.json().get("uid")
    for v in list(session.cookies.jar._cookies.values())[0].values():
        for k, i in v.items():
            session.cookies.set(i.name, i.value)

    return session, uid


async def get_cookies(session):
    cookies = []
    for cookie_jar in session.cookies.jar:
        cookie = {}
        for name in (
                "version",
                "name",
                "value",
                "port",
                "port_specified",
                "domain",
                "domain_specified",
                "domain_initial_dot",
                "path",
                "path_specified",
                "secure",
                "expires",
                "discard",
                "comment",
                "comment_url",
        ):
            attr = getattr(cookie_jar, name)
            cookie.update({name: attr})
        cookies.append(cookie)
    return cookies


def add_title_data(title, name, data):
    title_run = title.add_run(name, style=STYLE)
    title_run.font.color.rgb = RGBColor(118, 113, 113)
    title_run.bold = True
    title_run.font.size = Pt(12)
    title_run = title.add_run(f": {data}", style=STYLE)
    title_run.italic = True
    title_run.font.size = Pt(12)


async def subects_names(session, referenceFilter, user_id, uid):
    response = await post(session, SUBECT_URL, {
        "group_id": user_id,
        "is_user_id": 1
    })
    response_2 = await post(session, SUBECT_URL, {
        "group_id": uid,
        "is_user_id": 1
    })
    names = []
    try:
        res = []
        for r in response.json():
            res.extend(r['items'] or [])
        for r in res:
            if r.get("id") in referenceFilter:
                names.append(r.get("keyword"))
        try:
            for r in response_2.json():
                res.extend(r['items'] or [])
            for r in res:
                if r.get("id") in referenceFilter:
                    if r.get("keyword") not in names:
                        names.append(r.get("keyword"))
        except Exception:
            pass
        return names
    except Exception as e:
        try:
            from app import logger
        except Exception:
            logger = logging.getLogger('foo-logger')
        logger.error(f"subects_names {e}")
        return []


async def get_posts(session, thread_id, _from, _to, network_id, referenceFilter, friendly_ids=None, trustoption=None):
    if friendly_ids is None:
        friendly_ids = []
    limit = 200
    start = 0
    posts = []
    while True:
        payload = {
            "thread_id": thread_id,
            "from": _from,
            "to": _to,
            "limit": limit, "start": start, "sort": {"type": "date", "order": "desc", "name": "dateDown"},
            "friendly": friendly_ids,
            "trustoption": trustoption,
            "filter": {"network_id": network_id,
                       "referenceFilter": referenceFilter, "repostoption": "whatever"}
        }
        from resp import post

        response = await post(session, STATISTIC_POST_URL, payload)

        posts.extend(response.json().get("posts") or [])
        if not response.json().get("posts") or response.json().get("count") <= len(posts):
            break
        start += limit
    smi = 0
    friendly_smi = 0
    social = 0
    friendly_social = 0
    for post in posts:
        if post.get("network_id") == "4":
            smi += 1
            if check_friendly(post):
                friendly_smi += 1
        else:
            social += 1
            if check_friendly(post):
                friendly_social += 1

    return posts, smi, social, friendly_smi + friendly_social, friendly_smi, friendly_social, referenceFilter


def check_friendly(post):
    if post.get("friendly") in [2, 1, "2", "1"]:
        return True
    return False


async def get_posts_info(session, thread_id, periods_data, referenceFilter):
    table_gather = []
    for rec in referenceFilter:
        table_gather.append(get_posts(session, thread_id, periods_data.get("_from_data"), periods_data.get("_to_data"),
                                      NETWORK_IDS, [rec]))
    res = {}
    for (posts, smi, social, friendly, friendly_smi, friendly_social, rec) in await asyncio.gather(
            *table_gather):
        res[rec[0]] = (posts, smi, social, friendly, friendly_smi, friendly_social)
    return res


async def get_session_result(thread_id, _from, _to, referenceFilter, network_id, user_id, friendly_ids, trustoption):
    async with httpx.AsyncClient() as session:
        session, uid = await login(session)
        if uid is None:
            uid = user_id
        (posts, smi, social, friendly, friendly_smi, friendly_social, referenceFilter), names = await asyncio.gather(
            get_posts(session, thread_id, _from, _to, network_id, referenceFilter, friendly_ids, trustoption),
            subects_names(session, referenceFilter, user_id, uid)
        )
        return posts, smi, social, friendly, names


def insertHR(paragraph, line="double"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), line)
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def add_hyperlink_into_run(paragraph, run, url):
    runs = paragraph.runs
    for i in range(len(runs)):
        if runs[i].text == run.text:
            break
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    hyperlink.append(run._r)
    paragraph._p.insert(i, hyperlink)
    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)


def convert_date(date):
    try:
        return dateutil.parser.parse(date).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return dateutil.parser.parse(date).date().strftime("%d-%m-%Y")


async def docx_media(thread_id, _from, _to, referenceFilter, network_id, user_id, friendly_ids, trustoption=None, _sort=False):
    from app import UTC
    posts, smi, social, friendly, names = await get_session_result(thread_id, convert_date(_from), convert_date(_to),
                                                                   referenceFilter,
                                                                   network_id, user_id, friendly_ids, trustoption)
    document = Document()
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style(STYLE, WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.name = STYLE

    title = document.add_paragraph()
    title_run = title.add_run("СВОДКА ПУБЛИКАЦИЙ", style=STYLE)
    title_run.font.color.rgb = RGBColor(118, 113, 113)
    title_run.bold = True
    title_run.font.size = Pt(14)

    title = document.add_paragraph()

    if len(network_id) >= 2 and 4 in network_id:
        type_network = "СМИ и соцсети"
        number_networks = f"СМИ - {smi} публикаций, соцсети - {social} публикаций"
    elif 4 in network_id:
        type_network = "только СМИ"
        number_networks = f"СМИ - {smi} публикаций"
    else:
        type_network = "только соцсети"
        number_networks = f"соцсети - {social} публикаций"

    add_title_data(title, "Объекты", ", ".join(names))
    add_title_data(title, "\nИсточники публикаций", type_network)
    _from_str = dateutil.parser.parse(_from).strftime(DATE_FORMAT)
    _date_prepare = datetime.today() + timedelta(hours=UTC)
    if _date_prepare < dateutil.parser.parse(_to):
        _to_str = _date_prepare.strftime(DATE_FORMAT)
    else:
        _to_str = dateutil.parser.parse(_to).strftime(DATE_FORMAT)
    add_title_data(title, "\nВременной период", f"{_from_str} - {_to_str}")
    add_title_data(title, "\nДата подготовки отчета", _date_prepare.strftime(DATE_FORMAT))
    add_title_data(title, f"\nВсего сообщений", number_networks)
    if friendly % 10 == 1:
        friend_pud = f"{friendly} публикация"
    elif friendly % 10 == 2 or friendly % 10 == 3 or friendly % 10 == 4:
        friend_pud = f"{friendly} публикации"
    else:
        friend_pud = f"{friendly} публикаций"

    add_title_data(title, f"\nИз них дружественных ", friend_pud)

    insertHR(document.add_paragraph(), line="single")
    if _sort:
        posts.sort(key=lambda x: str(x['author']))
    for post in posts:
        post_paragraph = document.add_paragraph()
        paragraph_run = post_paragraph.add_run(post.get("author", ""), style=STYLE)
        paragraph_run.bold = True
        paragraph_run.font.size = Pt(12)
        paragraph_run = post_paragraph.add_run(f" {dateutil.parser.parse(post.get('created_date') or '').strftime(DATE_FORMAT)}"
                                                   "\n ", style=STYLE)
        paragraph_run.font.size = Pt(12)

        paragraph_ilnk = document.add_paragraph()
        url = post['uri'] or ''
        # if len(url) > 65:
         #     url = url[:65]
        paragraph_link = paragraph_ilnk.add_run(url, style=STYLE)
        paragraph_link.font.size = Pt(12)

        add_hyperlink_into_run(paragraph_ilnk, paragraph_link, post.get('uri') or '')

        if post['title']:
            post_paragraph_title = document.add_paragraph()
            paragraph_title = post_paragraph_title.add_run(post.get('title', ''), style=STYLE)
            paragraph_title.bold = True
            paragraph_title.font.size = Pt(12)
        if post.get("friendly"):
            post_paragraph_friendly = document.add_paragraph()

            post_paragraph_friendly = post_paragraph_friendly.add_run("Дружественный источник!", style=STYLE)
            post_paragraph_friendly.font.size = Pt(12)
            post_paragraph_friendly.bold = True
            post_paragraph_friendly.italic = True

        post_paragraph_text = document.add_paragraph()
        cleantext = re.sub(CLEANR, '', post.get('text') or '')
        post_paragraph_text = post_paragraph_text.add_run(cleantext, style=STYLE)
        post_paragraph_text.font.size = Pt(12)

        hr = document.add_paragraph()
        hr.style.font.size = Pt(1)

        insertHR(hr)
    return document
